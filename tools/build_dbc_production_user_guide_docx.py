from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "tools"))

import build_dbc_testing_methodology_docx as base  # noqa: E402


OUT = ROOT / "docs" / "DBC_Production_User_Guide_RU.docx"
SHOTS = ROOT / "docs" / "screenshots-guide"


def add_screenshot(doc, filename, caption):
    image_path = SHOTS / filename
    if not image_path.exists():
        base.add_callout(doc, "Скриншот не найден", str(image_path), base.AMBER)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.35))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(10)
    run = caption_p.add_run(caption)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(90, 105, 115)


def build():
    doc = Document()
    base.setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Dildin Build Control")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(base.BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Production User Guide: как из ТЗ получить проверенный результат через Guided Run")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(base.DARK_BLUE)

    base.add_callout(
        doc,
        "Главная идея",
        "DBC должен вести пользователя по одному понятному маршруту: ТЗ -> TaskContract -> WorkSlice -> HarnessRun -> checks/review/security -> EvidencePack -> Accept/Rework/Reject. Экспертные экраны остаются, но основной путь теперь находится в Guided Run.",
    )

    base.add_heading(doc, "1. Быстрый старт", 1)
    base.add_numbered(
        doc,
        [
            "Откройте DBC и начните с Control Tower. Там видно текущее состояние проекта и следующую рекомендуемую операцию.",
            "Перейдите в Guided Run. Это основной пользовательский сценарий для работы с ТЗ.",
            "Вставьте ТЗ, название задачи, acceptance criteria, allowed paths и forbidden paths.",
            "Нажмите Create and start safe run. DBC создаст задачу, контракт, work slice и HarnessRun в безопасном режиме.",
            "Нажимайте Advance run, пока работа не дойдёт до evidence_ready или не попросит approval.",
            "Соберите EvidencePack.",
            "Откройте Reports и примите решение Accept, Request rework или Reject.",
        ],
    )

    base.add_heading(doc, "2. Control Tower", 1)
    base.add_body(
        doc,
        "Control Tower показывает состояние проекта, pending approvals, текущий workflow и одну главную кнопку следующего действия. Если вы не знаете, куда идти дальше, начинайте отсюда.",
    )
    add_screenshot(doc, "01-control-tower.png", "Control Tower: операторская панель и Current Workflow.")

    base.add_heading(doc, "3. Guided Run: основной рабочий экран", 1)
    base.add_body(
        doc,
        "Guided Run связывает все внутренние сущности DBC в один сценарий. Пользователю не нужно отдельно понимать SQLite, .dbc, contracts, slices и loop compatibility layer: экран показывает шаги, блокировки и следующую команду.",
    )
    base.add_table(
        doc,
        ["Шаг", "Что делает пользователь", "Что делает DBC"],
        [
            ["TZ", "Вставляет ТЗ и acceptance criteria.", "Создаёт Task с границами scope и stop conditions."],
            ["Contract", "Проверяет смысл задачи.", "Создаёт/замораживает/approve TaskContract для safe path."],
            ["Slice", "Работает маленькими частями.", "Создаёт WorkSlice с allowed paths и commands."],
            ["Run", "Нажимает Advance run.", "Ведёт HarnessRun поверх существующего loop engine."],
            ["Evidence", "Генерирует EvidencePack.", "Связывает contract, slice, run, logs, review, security и approvals."],
            ["Decision", "Выбирает Accept/Rework/Reject.", "Сохраняет final decision через Harness Engine."],
        ],
        [1.1, 2.35, 3.05],
    )
    add_screenshot(doc, "02-guided-run.png", "Guided Run: единый путь от ТЗ до финального решения.")

    base.add_heading(doc, "4. Как заполнять ТЗ", 1)
    base.add_bullets(
        doc,
        [
            "Product task title: коротко назовите пользовательский результат.",
            "TZ / request: вставьте исходное ТЗ, контекст, ограничения и ожидаемое поведение.",
            "Acceptance criteria: перечислите условия, без которых работу нельзя принять.",
            "Allowed paths: укажите директории и файлы, где агент может работать.",
            "Out of scope / forbidden paths: добавьте .env, node_modules, target, dist и любые секретные зоны.",
            "Начинайте с safe/mock режима. Real provider включайте только через approval gate.",
        ],
    )
    base.add_callout(
        doc,
        "Правило безопасности",
        "DBC работает через официальные CLI/API/local terminal surfaces. Не используйте consumer/web интерфейсы через scraping. Real provider execution остаётся заблокированным без human approval.",
        base.AMBER,
    )

    base.add_heading(doc, "5. Approvals и блокировки", 1)
    base.add_body(
        doc,
        "Approval Queue нужна не для бюрократии, а как человеческий тормоз. Она должна явно показывать Spec, Plan, WorkSlice, Command, Real Provider и Evidence gates.",
    )
    base.add_table(
        doc,
        ["Gate", "Когда нужен approval"],
        [
            ["Spec approval", "Перед использованием замороженного TaskContract."],
            ["Plan approval", "Когда план влияет на scope, риск или стоимость."],
            ["WorkSlice approval", "Перед исполнением ограниченного куска работы."],
            ["Command approval", "Когда command policy классифицирует команду как approval_required."],
            ["Real provider approval", "Перед любым реальным запуском Codex/Claude."],
            ["Evidence acceptance", "Перед финальным Accept."],
        ],
        [1.75, 4.75],
    )

    base.add_heading(doc, "6. Reports и EvidencePack", 1)
    base.add_body(
        doc,
        "Reports отвечает на вопрос: можно ли принимать работу. Acceptance Checklist показывает, почему Accept доступен или заблокирован.",
    )
    base.add_table(
        doc,
        ["Решение", "Когда использовать"],
        [
            ["Accept", "Все loop steps прошли, approvals закрыты, EvidencePack создан, review/security не нашли blocker."],
            ["Request rework", "Работа близка, но есть missing tests, неполное evidence, UX issue или риск."],
            ["Reject", "Решение нарушает contract, command policy, forbidden paths или не достигает цели."],
        ],
        [1.4, 5.1],
    )
    add_screenshot(doc, "03-reports-checklist.png", "Reports: Acceptance Decision, Acceptance Checklist и Evidence Packs.")

    base.add_heading(doc, "7. Settings Quick Setup", 1)
    base.add_body(
        doc,
        "Settings разделён на быстрый слой и advanced-настройки. В начале смотрите Quick Setup: Codex, Claude, Local runner, real providers и safe mock.",
    )
    base.add_bullets(
        doc,
        [
            "Codex CLI: используйте exact path, если app PATH не видит команду.",
            "Claude Code: проверьте args template и prompt mode.",
            "Local runner: нужен для build/test/doctor/audit команд.",
            "Safe mock: рекомендуемый режим для обычной разработки и демонстрации.",
            "Real micro: включайте только после operator approval и clean checks.",
        ],
    )
    add_screenshot(doc, "04-settings-quick-setup.png", "Settings: Quick Setup и provider readiness.")

    base.add_heading(doc, "8. Команды проверки", 1)
    base.add_code(
        doc,
        "pnpm build\ncargo test --manifest-path src-tauri/Cargo.toml\npnpm launch-doctor\npnpm system-audit",
    )
    base.add_body(
        doc,
        "Production-ready состояние для текущей версии: build и cargo tests проходят, launch-doctor/system-audit не имеют blockers. Warnings допустимы, если они требуют human approval и понятны оператору.",
    )

    base.add_heading(doc, "9. Мини-чеклист перед работой", 1)
    base.add_bullets(
        doc,
        [
            "Пользователь понимает, что основной путь находится в Guided Run.",
            "CLI providers проверены или включён safe mock.",
            "Allowed/forbidden paths заданы до запуска.",
            "Real provider не запускается без approval.",
            "EvidencePack создан до Accept.",
            "Final decision содержит заметку, если выбран Rework или Reject.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Dildin Build Control · Production User Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
