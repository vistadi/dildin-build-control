from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "DBC_Testing_And_Methodology_Guide_RU.docx"
SHOTS = ROOT / "docs" / "screenshots-guide"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "DDEFE8"
AMBER = "FFF4D6"
BORDER = "D9E2EF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)


def set_cell_text(cell, text, bold=False, color=None, size=10.2):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    set_cell_border(cell)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles[f"Heading {level}"]
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 12)
    run.font.color.rgb = RGBColor.from_string(BLUE if level in (1, 2) else DARK_BLUE)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        tail = paragraph.add_run(text[len(bold_prefix):])
        tail.font.name = "Calibri"
        tail.font.size = Pt(11)
    else:
        run = paragraph.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.15
    for index, line in enumerate(text.splitlines()):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if index < len(text.splitlines()) - 1:
            paragraph.add_run("\n")
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "B7C9DF")
    set_cell_margins(cell, 110, 110, 160, 160)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.18
    run = paragraph.add_run(body)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE, size=10)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.8)
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT_GRAY)
    doc.add_paragraph()
    return table


def add_screenshot(doc, filename, caption):
    image_path = SHOTS / filename
    if not image_path.exists():
        add_callout(doc, "Скриншот не найден", str(image_path), AMBER)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.35))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(10)
    run = caption_p.add_run(caption)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(90, 105, 115)


def add_page_break(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def build():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Dildin Build Control")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Тест-план и инструкция: как из ТЗ довести задачу до evidence-backed результата")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    add_callout(
        doc,
        "Как пользоваться этим документом",
        "Сначала пройдите раздел 1 как тест-план приложения. Затем используйте раздел 2 как рабочую инструкцию: вставить ТЗ, создать задачу, запустить safe run, пройти approvals, собрать Evidence Pack и принять решение Accept/Rework/Reject.",
    )

    add_heading(doc, "1. Пошаговый план тестирования приложения", 1)
    add_body(doc, "Цель тестирования: убедиться, что DBC управляет задачей как local-first control tower, не запускает real providers без человека и сохраняет evidence в SQLite и .dbc artifacts.")

    add_heading(doc, "1.1 Быстрый smoke за 20 минут", 2)
    add_numbered(
        doc,
        [
            "Запустить приложение в desktop runtime: pnpm tauri dev. Для визуальной проверки UI можно использовать pnpm dev.",
            "Открыть Settings и проверить, что Codex CLI, Claude Code и Local Terminal не имеют blocker-ошибок.",
            "Нажать Test CLI и Check contract для Codex и Claude.",
            "Открыть Tasks и нажать Controlled smoke или Run smoke loop.",
            "Открыть Loops и убедиться, что появился HarnessRun или понятный безопасный статус ожидания.",
            "Открыть Approvals и проверить, что опасные действия требуют human approval.",
            "Открыть Reports и убедиться, что Evidence Pack или report-сводка доступны.",
            "Запустить команды проверки из терминала и убедиться, что нет blockers.",
        ],
    )
    add_code(doc, "pnpm build\ncargo test --manifest-path src-tauri/Cargo.toml\npnpm launch-doctor\npnpm system-audit")

    add_heading(doc, "1.2 Полный regression-план", 2)
    add_table(
        doc,
        ["Блок", "Что проверить", "Ожидаемый результат"],
        [
            ["Launch", "pnpm dev и pnpm tauri dev запускают UI.", "Открывается DBC shell без белого экрана и console errors."],
            ["Settings", "CLI discovery, exact path, Test CLI, Check contract.", "Provider status не Failed; PATH-warning допустим, blocker нет."],
            ["Task intake", "Создать задачу из ТЗ: title, TZ, criteria, paths, risk, budget.", "Задача появляется в списке Tasks."],
            ["Preflight", "Открыть Preflight для задачи.", "Gates показывают ready/warnings/blockers с понятным текстом."],
            ["Safe run", "Нажать Start safe run.", "Создаются TaskContract, WorkSlice, HarnessRun; legacy loop не ломается."],
            ["Approvals", "Проверить Spec/Plan/Slice/Command/Real Provider/Evidence gates.", "Real provider execution заблокирован без human approval."],
            ["Loops", "Нажимать Advance до self-check/review/security/evidence_ready.", "Статусы движутся по Harness lifecycle, ошибки видны в карточке run."],
            ["Evidence", "Нажать Evidence Pack.", "Создаётся manifest со ссылками на contract, slice, run journal, reports."],
            ["Reports", "Проверить final decision.", "Можно принять Accept, Rework или Reject на основе evidence."],
            ["Persistence", "Перезапустить приложение.", "Проект, задачи, runs и artifacts восстанавливаются из SQLite/.dbc."],
        ],
        [1.2, 2.55, 2.75],
    )

    add_heading(doc, "1.3 Что считать успешным тестом", 2)
    add_bullets(
        doc,
        [
            "Пользователь может начать с обычного ТЗ, не зная внутренних сущностей Harness.",
            "Основной путь работает через Start safe run, advanced controls не обязательны.",
            "Ни один real provider не запускается без явного человеческого approval.",
            "Каждый итог имеет evidence: contract, slice, run journal, checks, reports и final decision.",
            "После перезапуска приложения данные не пропадают.",
            "Команды pnpm build, cargo test, launch-doctor и system-audit проходят без blockers.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "2. Инструкция: как из ТЗ сделать готовый продукт", 1)
    add_body(doc, "Рабочий маршрут DBC выглядит так:")
    add_code(doc, "ТЗ -> Task -> Spec/Contract -> Approval -> WorkSlice -> HarnessRun -> Self-check -> Review -> Security -> Evidence Pack -> Accept/Rework/Reject -> Ship")

    add_heading(doc, "Шаг 1. Проверить Control Tower", 2)
    add_body(doc, "Откройте Control Tower. Здесь видно общее состояние проекта: open tasks, pending approvals, cost events, launch doctor и real micro readiness.")
    add_screenshot(doc, "01-control-tower.png", "Control Tower: общая панель здоровья проекта и запуск provider loop.")

    add_heading(doc, "Шаг 2. Перенести ТЗ в Task Composer", 2)
    add_body(doc, "Откройте Tasks и заполните поля так, чтобы задача стала исполнимым контрактом:")
    add_bullets(
        doc,
        [
            "Task title: короткое имя результата.",
            "Free-form TZ: исходное ТЗ, контекст, ограничения и риски.",
            "Acceptance criteria: что должно быть истинно, чтобы принять работу.",
            "Constraints: правила, которые нельзя нарушать.",
            "Affected paths и Allowed paths: где можно работать.",
            "Denied paths: .env, node_modules, target/build artifacts и другие запретные зоны.",
            "Required reviewers: Reviewer, Security, Product Owner.",
            "Stop conditions: когда DBC должен остановиться и попросить человека.",
        ],
    )
    add_screenshot(doc, "02-tasks-start-safe-run.png", "Tasks: Task Composer, smoke-controls и основной путь через Start safe run.")

    add_heading(doc, "Шаг 3. Запустить безопасный путь", 2)
    add_body(doc, "Для обычной работы нажмите Start safe run на карточке задачи. DBC должен сам выполнить техническую рутину: создать TaskContract, заморозить spec, создать WorkSlice, связать HarnessRun и запустить compatibility loop.")
    add_callout(
        doc,
        "Не начинайте с Advanced controls",
        "Create Contract, Freeze Spec, Approve Contract, Create Slice и Start Harness нужны для ручной отладки. Пользовательский путь должен идти через Start safe run.",
        AMBER,
    )

    add_heading(doc, "Шаг 4. Проверить Preflight", 2)
    add_body(doc, "Preflight показывает, готова ли задача к запуску. Если есть blockers, их нужно исправить до запуска real provider. Warnings допустимы только если вы понимаете риск.")
    add_screenshot(doc, "03-preflight-gates.png", "Preflight: gates по задаче, проекту, provider strategy, command policy и approvals.")

    add_heading(doc, "Шаг 5. Вести работу через Loops", 2)
    add_body(doc, "В Loops смотрите Harness Runs. Здесь должны быть связаны Contract, Slice, compatibility loop и Evidence Pack. Нажимайте Advance, пока run не дойдёт до evidence_ready или не попросит approval/rework.")
    add_table(
        doc,
        ["Статус", "Что означает"],
        [
            ["planned", "План создан, но работа ещё не запущена."],
            ["waiting_approval", "DBC ждёт человека."],
            ["slice_running", "Исполняется ограниченный WorkSlice."],
            ["self_checked", "Первичная самопроверка завершена."],
            ["reviewed", "Проверка reviewer-ролью завершена."],
            ["security_reviewed", "Security review завершён."],
            ["evidence_ready", "Можно собрать Evidence Pack."],
            ["accepted/rework/rejected", "Финальное решение по результату."],
        ],
        [1.65, 4.85],
    )
    add_screenshot(doc, "04-loops-harness-run.png", "Loops: Harness Runs, Loop History, Run Journal и кнопки Advance/Evidence Pack.")

    add_heading(doc, "Шаг 6. Проверить human approvals", 2)
    add_body(doc, "Откройте Approvals. Здесь должны появляться gates по фазам Spec, Plan, WorkSlice, Command, Real Provider и Evidence. Особенно важно: real provider execution должен оставаться заблокированным без ручного подтверждения.")
    add_screenshot(doc, "05-approvals-queue.png", "Approvals: очередь человеческих решений и sensitive gates.")

    add_heading(doc, "Шаг 7. Собрать Evidence Pack и принять решение", 2)
    add_body(doc, "Когда run дошёл до evidence_ready, нажмите Evidence Pack в Loops, затем откройте Reports. Решение принимается не по ощущению, а по evidence.")
    add_table(
        doc,
        ["Решение", "Когда выбирать"],
        [
            ["Accept", "Acceptance criteria выполнены, tests/build прошли, security risks закрыты."],
            ["Rework", "Идея верная, но есть missing tests, UX issue, risk или неполное evidence."],
            ["Reject", "Решение нарушает contract, policy, forbidden paths или не достигает бизнес-цели."],
        ],
        [1.25, 5.25],
    )
    add_screenshot(doc, "06-reports-evidence.png", "Reports: Evidence Packs, structured reports и итоговое acceptance/rework решение.")

    add_heading(doc, "Шаг 8. Настроить CLI и режимы", 2)
    add_body(doc, "В Settings хранится practical control plane: CLI paths, provider contracts, command policy, local terminal, profiles mock/real-micro.")
    add_table(
        doc,
        ["Provider", "Как использовать"],
        [
            ["Codex CLI", "Developer/Team Lead: планирование, код, тесты. Используйте exact path и официальный CLI."],
            ["Claude Code", "Reviewer/QA/Security: review diff, logs, security notes. Используйте exact path и -p/stdin contract."],
            ["Local Terminal", "DevOps: pnpm build, cargo test, launch-doctor, system-audit по command policy."],
        ],
        [1.45, 5.05],
    )
    add_screenshot(doc, "07-settings-cli.png", "Settings: CLI discovery, exact path, run modes, provider contracts и command policy.")

    add_page_break(doc)
    add_heading(doc, "3. Команды для терминала", 1)
    add_table(
        doc,
        ["Команда", "Для чего"],
        [
            ["pnpm dev", "Открыть web-preview UI для визуальной проверки."],
            ["pnpm tauri dev", "Запустить настоящий desktop runtime с Tauri backend."],
            ["pnpm build", "Проверить TypeScript/Vite сборку."],
            ["cargo test --manifest-path src-tauri/Cargo.toml", "Проверить Rust backend."],
            ["pnpm controlled-smoke", "Проверить безопасный smoke loop."],
            ["pnpm launch-doctor", "Проверить readiness и blockers."],
            ["pnpm system-audit", "Проверить целостность системы."],
            ["pnpm provider-contracts", "Проверить CLI/provider contracts."],
            ["pnpm approval-queue", "Проверить очередь approvals."],
        ],
        [2.45, 4.05],
    )

    add_heading(doc, "4. Мини-чеклист перед real provider", 1)
    add_bullets(
        doc,
        [
            "Codex CLI и Claude Code настроены через exact path, а не только через PATH command.",
            "Run mode остаётся mock, пока smoke и audits не зелёные.",
            "Approval Queue не содержит незакрытых blockers.",
            "Command policy не ослаблена: опасные команды требуют approval или deny.",
            ".env и секреты входят в denied paths и не попадают в logs/.dbc artifacts.",
            "Есть понятный rollback/rework путь.",
        ],
    )
    add_callout(
        doc,
        "Главное правило безопасности",
        "DBC должен автоматизировать официальные CLI/API/local terminal surfaces. Не используйте consumer/web-интерфейсы через scraping. Real provider запускается только после human approval.",
        AMBER,
    )

    add_heading(doc, "5. Что делать, если тест не проходит", 1)
    add_table(
        doc,
        ["Симптом", "Действие"],
        [
            ["CLI not found", "В Settings вставить exact path из which codex / which claude и нажать Test CLI."],
            ["stdin is not a terminal", "Проверить args template: Codex через exec ..., Claude через -p/stdin."],
            ["HarnessRun не появился", "Открыть Tasks, убедиться, что задача сохранена, затем Start safe run; проверить .dbc и Loops Refresh."],
            ["Evidence Pack неактивен", "В Loops нажимать Advance до evidence_ready или устранить blocker/rework."],
            ["launch-doctor warnings", "Warnings допустимы для preview; blockers нужно исправить."],
            ["Real provider blocked", "Это нормальная защита. Нужно approval gate и готовый operator checklist."],
        ],
        [2.05, 4.45],
    )

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Dildin Build Control · Testing & Methodology Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
