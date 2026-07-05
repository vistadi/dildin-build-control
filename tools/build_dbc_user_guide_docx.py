from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "DBC_User_Guide_RU.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    style = doc.styles[f"Heading {level}"]
    paragraph.style = style
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(BLUE)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string(BLUE)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    for line in text.splitlines():
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
        if line != text.splitlines()[-1]:
            paragraph.add_run("\n")
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, "B7C9DF")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(body)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT_GRAY)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Dildin Build Control")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Короткая инструкция пользователя: простой режим, CLI-настройка и безопасный запуск")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    add_callout(
        doc,
        "Главная идея",
        "Обычный пользователь не должен вручную нажимать Create Contract, Freeze Spec, Approve Contract, Create Slice и Start Harness. "
        "В нормальной работе используется одна кнопка: Start safe run.",
    )

    add_heading(doc, "1. Самый простой сценарий", 1)
    add_code(doc, "Tasks -> Save task -> Start safe run -> Loops -> Advance -> Evidence Pack -> Reports")
    add_table(
        doc,
        ["Шаг", "Что сделать", "Что получится"],
        [
            ["1", "Открыть Tasks и описать задачу в Task Composer.", "Появится задача в списке."],
            ["2", "Нажать Save task.", "Задача сохранится в проекте."],
            ["3", "На карточке задачи нажать Start safe run.", "DBC сам создаст contract, slice и HarnessRun."],
            ["4", "Открыть Loops и нажимать Advance.", "Система пройдет self-check, review и security stages."],
            ["5", "Когда статус станет evidence_ready, нажать Evidence Pack.", "Будет создан пакет доказательств."],
            ["6", "Открыть Reports.", "Можно проверить итоговый отчет и принять решение."],
        ],
        [0.55, 2.65, 3.3],
    )

    add_heading(doc, "2. Что не трогать в обычной работе", 1)
    add_body(doc, "Эти кнопки нужны только для ручной отладки или advanced-сценариев:")
    add_bullets(
        doc,
        [
            "Create Contract",
            "Freeze Spec",
            "Approve Contract",
            "Create Slice",
            "Approve Slice",
            "Start Harness",
            "Start legacy loop",
        ],
    )
    add_body(doc, "Они спрятаны в Advanced controls. В обычной работе достаточно Start safe run.")

    add_heading(doc, "3. Один раз настроить Codex и Claude", 1)
    add_numbered(
        doc,
        [
            "Открыть Settings.",
            "Для Codex CLI указать exact path из команды which codex.",
            "Для Claude Code указать exact path из команды which claude.",
            "Нажать Test CLI для каждого provider.",
            "Нажать Check contract для каждого provider.",
            "Нажать Sync .dbc config.",
        ],
    )
    add_table(
        doc,
        ["Provider", "Args template", "Prompt mode", "Run mode"],
        [
            ["Codex CLI", 'exec --skip-git-repo-check --sandbox workspace-write --cd "{{cwd}}"', "stdin", "mock"],
            ["Claude Code", "-p", "stdin", "mock"],
        ],
        [1.15, 3.35, 1.0, 1.0],
    )

    add_heading(doc, "4. Перед каждой работой", 1)
    add_numbered(
        doc,
        [
            "Открыть Settings.",
            "Нажать Load .dbc config.",
            "Проверить, что Codex и Claude не красные.",
            "Оставить Run mode: mock, если не нужно тратить реальные provider calls.",
            "Открыть Tasks и работать через Start safe run.",
        ],
    )

    add_heading(doc, "5. Безопасная проверка системы", 1)
    add_body(doc, "Перед реальной задачей лучше сделать controlled smoke.")
    add_code(doc, "pnpm controlled-smoke\npnpm launch-doctor\npnpm system-audit")
    add_body(doc, "Нормальный безопасный статус:")
    add_code(doc, "ready_for_human_approval")
    add_body(doc, "Это не ошибка. Это означает, что DBC специально ждет человека перед запуском real providers.")

    add_heading(doc, "6. Когда можно включать real provider", 1)
    add_callout(
        doc,
        "Правило",
        "Не включать Run mode: real, пока не прошли controlled smoke, provider-contracts, approval-queue, launch-doctor и system-audit.",
    )
    add_code(doc, "pnpm controlled-smoke\npnpm provider-contracts\npnpm approval-queue\npnpm launch-doctor\npnpm system-audit")
    add_body(doc, "Если цель: Codex пишет код, Claude проверяет, используйте такую схему ролей:")
    add_table(
        doc,
        ["Роль", "Provider"],
        [
            ["Developer / Team Lead", "Codex CLI"],
            ["Reviewer / QA / Security", "Claude Code"],
            ["DevOps", "Local Terminal"],
        ],
        [2.2, 4.3],
    )

    add_heading(doc, "7. Где смотреть результат", 1)
    add_table(
        doc,
        ["Экран", "Для чего"],
        [
            ["Tasks", "Создать задачу и нажать Start safe run."],
            ["Loops", "Нажимать Advance и смотреть HarnessRun status."],
            ["Approvals", "Проверять human gates и real provider approvals."],
            ["Reports", "Смотреть EvidencePack и финальный отчет."],
            ["Settings", "Настраивать Codex, Claude, local terminal и command policy."],
        ],
        [1.35, 5.15],
    )

    add_heading(doc, "8. Нормальные и плохие статусы", 1)
    add_table(
        doc,
        ["Нормально", "Плохо"],
        [
            ["running", "blocked"],
            ["slice_running", "failed"],
            ["self_checked", "rejected"],
            ["reviewed", ""],
            ["security_reviewed", ""],
            ["evidence_ready", ""],
            ["accepted", ""],
            ["ready_for_human_approval", ""],
        ],
        [3.25, 3.25],
    )

    add_heading(doc, "9. Если что-то не работает", 1)
    add_table(
        doc,
        ["Проблема", "Что сделать"],
        [
            ["CLI not found", "Вставить exact path из which codex или which claude."],
            ["stdin is not a terminal", "Проверить args: Codex должен запускаться через exec ..., Claude через -p."],
            ["Start safe run не сработал", "Открыть audit/Loops и посмотреть последнее сообщение ошибки."],
            ["Evidence Pack неактивен", "В Loops нажимать Advance, пока статус не станет evidence_ready."],
            ["launch-doctor показывает warnings", "Смотреть blockers. Если blockers: 0, это не критично."],
        ],
        [2.05, 4.45],
    )

    add_heading(doc, "10. Где лежат файлы", 1)
    add_table(
        doc,
        ["Папка", "Что внутри"],
        [
            [".dbc/contracts/", "Frozen specs / TaskContract JSON."],
            [".dbc/slices/", "WorkSlice JSON."],
            [".dbc/harness-runs/", "HarnessRun manifests and events."],
            [".dbc/packs/", "EvidencePack manifests."],
            [".dbc/reports/", "Acceptance reports."],
            [".dbc/approval-gates/", "Harness gates."],
            [".dbc/approval-queue/", "Unified approval queue."],
        ],
        [2.05, 4.45],
    )

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Dildin Build Control · User Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
